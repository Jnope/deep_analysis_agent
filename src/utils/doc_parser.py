import os
from typing import List

from loguru import logger


def parse_file(file_path: str) -> str:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        ".txt": _parse_txt,
        ".md": _parse_txt,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".doc": _parse_docx,
        ".html": _parse_html,
        ".htm": _parse_html,
    }

    parser = parsers.get(ext)
    if parser is None:
        logger.warning(f"不支持的文件类型: {ext}，按纯文本处理")
        return _parse_txt(file_path)

    logger.info(f"解析文件: {file_path} (类型: {ext})")
    return parser(file_path)


def parse_file_chunked(file_path: str, chunk_size: int = 20000) -> List[str]:
    content = parse_file(file_path)
    if len(content) <= chunk_size:
        return [content] if content else []

    chunks = []
    for i in range(0, len(content), chunk_size):
        chunks.append(content[i:i + chunk_size])
    logger.info(f"文件分块: {len(chunks)} 块 (每块约 {chunk_size} 字符)")
    return chunks


def _parse_txt(file_path: str) -> str:
    for encoding in ["utf-8", "gbk", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    logger.error(f"无法解码文件: {file_path}")
    return ""


def _parse_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- 第 {i + 1} 页 ---\n{text}")
    logger.info(f"PDF 共 {len(reader.pages)} 页，提取到 {len(pages)} 页有效内容")
    return "\n\n".join(pages)


def _parse_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))

    parts = []
    if paragraphs:
        parts.append("\n\n".join(paragraphs))
    if tables:
        parts.append("--- 表格 ---\n" + "\n\n".join(tables))

    logger.info(f"DOCX: {len(paragraphs)} 段落, {len(tables)} 个表格")
    return "\n\n".join(parts)


def _parse_html(file_path: str) -> str:
    from bs4 import BeautifulSoup

    for encoding in ["utf-8", "gbk", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                html = f.read()
            break
        except UnicodeDecodeError:
            continue
    else:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    lines = [line for line in text.splitlines() if line.strip()]
    return "\n".join(lines)